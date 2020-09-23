# @Time    : 2020/9/23 9:15
# @Author  : Arvin
# @File    : DocumentOp.py
# @Software: PyCharm
# @Title   :

import os
import docx
from win32com import client
import re
import xlwt

if __name__ == '__main__':
    path = 'E:/Arvin/Document/WebCrawler/DocumenOp/documents'
    files = os.listdir(path)
    word = client.Dispatch("Word.Application")
    print(files)
    names = []
    locations = []
    times = []
    dates = []
    salaries = []
    for file in files:
        try:
            print(file)
            Name = re.search('[_][\u4e00-\u9fa5]{1,3}[_]', file)
            name = Name.group(0)
            print(name[1:4])
            names.append(name[1:4])
            file_path = 'E:/Arvin/Document/WebCrawler/DocumenOp/documents/' + file
            print(file_path)
            doc = word.Documents.Open(file_path)
            file_path2 = 'E:/Arvin/Document/WebCrawler/DocumenOp/documents2/' + file
            doc.SaveAs(file_path2+'x', 16)
            doc.Close()
            file_doc = docx.Document(file_path2+'x')
            text = file_doc.tables[1].rows[1].cells[0].text
            time = re.search('\d{1,}[\u4e00-\u9fa5]{4,}', text)
            try:
                times.append(time.group(0))
                print(time.group(0))
            except:
                times.append(0)
            location = re.search('[\u4e00-\u9fa5]{4,}.[\u4e00-\u9fa5]{1,}', text)
            try:
                locations.append(location.group(0))
                print(location.group(0))
            except:
                locations.append(0)
            date = re.search('[\u4e00-\u9fa5]{5,}.{1,}', file_doc.paragraphs[0].text)
            try:
                dates.append(date.group(0))
                print(date.group(0))
            except:
                dates.append(0)
            try:
                salay_t = re.search('\d+.\d+', file_doc.tables[2].rows[0].cells[0].text)
                salay_i = salay_t.group(0).split('-')
                if len(salay_i)==1:
                    salaries.append((int)(salay_i[0]))
                    print(salay_i[0])
                else:
                    salay =  ((int)(salay_i[0]) + ((int)(salay_i[1])))/2
                    print(salay)
                    salaries.append(salay)
            except:
                salaries.append(0)
                print(0)
        except:
            print(file, 'error')
    word.Quit()
    wb = xlwt.Workbook()
    sh1 = wb.add_sheet('sheet1')
    for num in range(0, len(names)):
        sh1.write(num, 0, names[num])
        sh1.write(num, 1, locations[num])
        sh1.write(num, 2, times[num])
        sh1.write(num, 3, dates[num])
        sh1.write(num, 4, salaries[num])
        print(num)
    wb.save('test_w_2.xls')
