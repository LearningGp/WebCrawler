# @Time    : 2020/9/23 10:34
# @Author  : Arvin
# @File    : DocumentOp2.py
# @Software: PyCharm
# @Title   :

import os
import docx
from win32com import client
import re

if __name__ == '__main__':
    path = 'E:/Arvin/Document/WebCrawler/DocumenOp/documents'
    files = os.listdir(path)
    word = client.Dispatch("Word.Application")
    print(files)
    for file in files:
        print(file)
        file_path = 'E:/Arvin/Document/WebCrawler/DocumenOp/documents/' + file
        print(file_path)
        doc = word.Documents.Open(file_path)
        doc.SaveAs(file_path+'x', 16)
        doc.Close()
        file_doc = docx.Document(file_path+'x')
        text = file_doc.tables[0].rows[0].cells[0].text
        print(text)
        for table in file_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    print(cell.text)
        # time = re.search('\d{1,}[\u4e00-\u9fa5]{4,}', text)
        # location = re.search('[\u4e00-\u9fa5]{4,}.[\u4e00-\u9fa5]{1,}',text)
        # print(location.group(0))
        # print(time.group(0))
        # date = re.search('[\u4e00-\u9fa5]{5,}.{1,}', file_doc.paragraphs[0].text)
        # print(date.group(0))
    word.Quit()